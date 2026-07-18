import sys
from datetime import UTC, datetime
from pathlib import Path
from typing import IO

# Ensure local packages are accessible
sys.path.insert(0, str(Path(__file__).parent.parent / "packages" / "content" / "src"))

from content.processors.docx import DOCXProcessor
from content.resource import (
    AbstractStreamReference,
    Checksum,
    ChecksumAlgorithm,
    ContentSource,
    LifecycleState,
    ResourceHandle,
    ResourceMetadata as HandleMetadata,
)


class LocalFileStreamReference(AbstractStreamReference):
    def __init__(self, file_path: Path) -> None:
        self.file_path = file_path

    def open_stream(self) -> IO[bytes]:
        return self.file_path.open("rb")


def create_sample_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.core_properties.title = "Sample DOCX"
    doc.add_heading("Welcome to DOCX Processor", 0)
    doc.add_paragraph("This is a sample document.")
    doc.add_paragraph("List item 1", style="List Bullet")
    doc.add_paragraph("List item 2", style="List Bullet")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    if len(sys.argv) > 1:
        if sys.argv[1] in ("-h", "--help"):
            print("Usage: uv run python dev/demo_docx_processor.py [path_to_docx]")
            print("If no path is provided, it defaults to:")
            print("  dev/sample_documents/sample.docx (generated on the fly if missing)")
            sys.exit(0)
        docx_path = Path(sys.argv[1]).resolve()
    else:
        docx_path = Path(__file__).parent / "sample_documents" / "sample.docx"

    if not docx_path.exists():
        if docx_path.name == "sample.docx":
            print("Generating sample.docx...")
            create_sample_docx(docx_path)
        else:
            print(f"Error: Could not find DOCX file at {docx_path}")
            sys.exit(1)

    stream_ref = LocalFileStreamReference(docx_path)

    handle = ResourceHandle(
        id="demo_docx_001",
        filename=docx_path.name,
        extension=docx_path.suffix,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source=ContentSource.UPLOAD,
        checksum=Checksum(algorithm=ChecksumAlgorithm.SHA256, value="dummy_checksum"),
        size_bytes=docx_path.stat().st_size,
        created_at=datetime.now(UTC),
        metadata=HandleMetadata(),
        stream_reference=stream_ref,
        lifecycle_state=LifecycleState.CREATED,
    )

    print(f"Processing: {docx_path.name}...\n")
    processor = DOCXProcessor()
    doc = processor.process(handle)

    print("-" * 40)
    print(f"Title       : {doc.title}")

    total_blocks = sum(len(page.blocks) for page in doc.pages)
    print(f"Block count : {total_blocks}")
    print(f"Statistics  : {doc.statistics}")
    print("-" * 40)
    print()

    if not doc.pages or not doc.pages[0].blocks:
        print("No blocks found.")
        return

    first_page = doc.pages[0]
    print("\nFirst 5 blocks text:")

    for i, block in enumerate(first_page.blocks[:5], start=1):
        text = block.text.replace("\n", " ")
        if len(text) > 100:
            text = text[:97] + "..."
        print(f"  [{i}] ({block.block_type.name}) {text}")


if __name__ == "__main__":
    main()
