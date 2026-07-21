import io
from collections.abc import Callable
from datetime import UTC, datetime
from typing import Any

import pytest
from docx import Document

from content.normalized.enums import BlockType
from content.processors.docx import DOCXProcessor
from content.processors.docx.exceptions import DOCXCorruptedError, DOCXEmptyError
from content.resource import (
    AbstractStreamReference,
    Checksum,
    ChecksumAlgorithm,
    ContentSource,
    LifecycleState,
    ResourceHandle,
    ResourceMetadata as HandleMetadata,
)


class MemoryStreamReference(AbstractStreamReference):
    def __init__(self, data: bytes) -> None:
        self.data = data

    def open_stream(self) -> io.BytesIO:
        return io.BytesIO(self.data)


def create_handle(data: bytes, filename: str = "test.docx") -> ResourceHandle:
    return ResourceHandle(
        id="test_id_123",
        filename=filename,
        extension=".docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source=ContentSource.UPLOAD,
        checksum=Checksum(algorithm=ChecksumAlgorithm.SHA256, value="dummy"),
        size_bytes=len(data),
        created_at=datetime.now(UTC),
        metadata=HandleMetadata(),
        stream_reference=MemoryStreamReference(data),
        lifecycle_state=LifecycleState.CREATED,
    )


def create_docx_bytes(build_fn: Callable[[Any], None] | None = None) -> bytes:
    doc = Document()
    if build_fn:
        build_fn(doc)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_docx_processor_success() -> None:
    def build(doc: Any) -> None:
        doc.core_properties.title = "Core Title"
        doc.add_heading("Heading 1", 1)
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Item 1", style="List Bullet")

    data = create_docx_bytes(build)
    handle = create_handle(data)
    processor = DOCXProcessor()

    doc = processor.process(handle)
    assert doc.title == "Core Title"

    blocks = doc.pages[0].blocks
    assert len(blocks) == 3
    assert blocks[0].block_type == BlockType.HEADING
    assert blocks[0].text == "Heading 1"

    assert blocks[1].block_type == BlockType.PARAGRAPH
    assert blocks[1].text == "Paragraph 1"

    assert blocks[2].block_type == BlockType.LIST
    assert blocks[2].text == "Item 1"


def test_docx_processor_table() -> None:
    def build(doc: Any) -> None:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

    data = create_docx_bytes(build)
    handle = create_handle(data)
    processor = DOCXProcessor()

    doc = processor.process(handle)
    blocks = doc.pages[0].blocks
    assert len(blocks) == 1
    assert blocks[0].block_type == BlockType.TABLE
    assert "| A1 | B1 |" in blocks[0].text
    assert "| A2 | B2 |" in blocks[0].text


def test_docx_processor_empty() -> None:
    data = create_docx_bytes()
    handle = create_handle(data)
    processor = DOCXProcessor()

    with pytest.raises(DOCXEmptyError):
        processor.process(handle)


def test_docx_processor_corrupted() -> None:
    handle = create_handle(b"not a docx zip file")
    processor = DOCXProcessor()

    with pytest.raises(DOCXCorruptedError):
        processor.process(handle)


def test_docx_processor_title_fallback_heading() -> None:
    def build(doc: Any) -> None:
        doc.add_heading("Fallback Heading", 1)
        doc.add_paragraph("Paragraph")

    data = create_docx_bytes(build)
    handle = create_handle(data)
    processor = DOCXProcessor()
    doc = processor.process(handle)

    assert doc.title == "Fallback Heading"


def test_docx_processor_title_fallback_filename() -> None:
    def build(doc: Any) -> None:
        # Add table only, no heading or paragraph
        doc.add_table(rows=1, cols=1)

    data = create_docx_bytes(build)
    handle = create_handle(data, filename="fallback.docx")
    processor = DOCXProcessor()
    doc = processor.process(handle)

    assert doc.title == "fallback.docx"


def test_plugin_registration() -> None:
    from content.plugins.registry import ProcessorRegistry

    registry = ProcessorRegistry()
    registry.register(DOCXProcessor())

    processor = registry.processor_for_extension("docx")
    assert isinstance(processor, DOCXProcessor)
